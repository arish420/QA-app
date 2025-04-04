import streamlit as st

# IMPORTANT: set_page_config must be the first st command in the script
st.set_page_config(page_title="Document Q&A Extractor", page_icon="📄", layout="wide")

import os
import tempfile
from pathlib import Path
import pytesseract
from pdf2image import convert_from_path
import pdfplumber 
import re
from docx import Document
import sys
import subprocess
import importlib 

# First try to import dotenv
try:
    from dotenv import load_dotenv
    # Load environment variables from .env file
    load_dotenv()
except ImportError:
    st.sidebar.warning("python-dotenv not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-dotenv"])
    from dotenv import load_dotenv
    load_dotenv()
    st.sidebar.success("python-dotenv installed successfully.")

# Initialize session state for API key
if 'api_key' not in st.session_state:
    st.session_state.api_key = None

# Store extracted questions in session state
if 'questions' not in st.session_state:
    st.session_state.questions = []

# Store generated answers in session state
if 'questions_answers' not in st.session_state:
    st.session_state.questions_answers = []

# Check if OpenAI is installed and install the correct version if needed
def check_openai_version():
    try:
        import openai
        current_version = openai.__version__
        st.sidebar.info(f"OpenAI version: {current_version}")
        if not current_version.startswith("0.28"):
            try:
                st.sidebar.warning("OpenAI version mismatch. Installing compatible version...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", "--force-reinstall", "openai==0.28.0"])
                st.sidebar.success("OpenAI 0.28.0 installed. Please refresh the page.")
                st.stop()
            except Exception as e:
                st.sidebar.error(f"Failed to install OpenAI 0.28.0: {str(e)}")
                st.sidebar.warning("Will attempt to use current version, but functionality may be limited.")
    except ImportError:
        try:
            st.sidebar.warning("OpenAI not found. Installing compatible version...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openai==0.28.0"])
            st.sidebar.success("OpenAI 0.28.0 installed. Please refresh the page.")
            st.stop()
        except Exception as e:
            st.sidebar.error(f"Failed to install OpenAI: {str(e)}")
            st.stop()
    
    # Import after ensuring the correct version
    import openai
    return openai

# Set Tesseract path based on operating system
def setup_tesseract():
    tesseract_found = False
    if os.name == "nt":  # Windows
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            tesseract_found = True
        else:
            st.sidebar.warning("⚠️ Tesseract OCR not found in default Windows location.")
            st.sidebar.markdown("""
            Please install Tesseract OCR:
            1. Download from: https://github.com/UB-Mannheim/tesseract/wiki
            2. Install and ensure it's in your PATH
            """)
    else:  # macOS or Linux
        try:
            # Try to find tesseract in the path
            tesseract_cmd = subprocess.run(["which", "tesseract"], 
                                           capture_output=True, 
                                           text=True, 
                                           check=True).stdout.strip()
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            tesseract_found = True
        except (subprocess.CalledProcessError, FileNotFoundError):
            # If not found, try default path
            default_path = "/usr/bin/tesseract"
            if os.path.exists(default_path):
                pytesseract.pytesseract.tesseract_cmd = default_path
                tesseract_found = True
            else:
                st.sidebar.warning("⚠️ Tesseract OCR not found. OCR functionality will not work.")
                if os.name == "posix" and "darwin" in sys.platform:  # macOS
                    st.sidebar.markdown("""
                    Please install Tesseract OCR:
                    ```bash
                    brew install tesseract
                    ```
                    """)
                else:  # Linux
                    st.sidebar.markdown("""
                    Please install Tesseract OCR:
                    ```bash
                    sudo apt-get install -y tesseract-ocr
                    ```
                    """)
    
    # Verify installation if tesseract is found
    if tesseract_found:
        try:
            version = subprocess.run([pytesseract.pytesseract.tesseract_cmd, "--version"], 
                                     capture_output=True, 
                                     text=True).stdout
            st.sidebar.success(f"✅ Tesseract OCR found: {version.split()[1] if version else ''}")
        except Exception:
            st.sidebar.warning("⚠️ Tesseract OCR found but not working properly.")
    
    return tesseract_found

# Functions for text extraction
def extract_text_with_ocr(pdf_path):
    try:
        # First check if tesseract is available
        if not hasattr(pytesseract.pytesseract, 'tesseract_cmd') or not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
            st.error("Tesseract OCR not available. Cannot process scanned document.")
            return ""
        
        try:
            images = convert_from_path(pdf_path)
        except Exception as e:
            if "poppler" in str(e).lower():
                st.error("Poppler not installed. Cannot process PDF with OCR.")
                st.info("For local installation: On Windows use conda install -c conda-forge poppler, on Linux use apt-get install poppler-utils")
                return ""
            else:
                raise e
            
        full_text = ""
        total_images = len(images)
        progress_placeholder = st.empty()
        
        for i, image in enumerate(images):
            full_text += pytesseract.image_to_string(image) + "\n"
            # Update progress safely
            percentage = int(100 * (i + 1) / total_images)
            progress_placeholder.text(f"OCR Processing: {percentage}% ({i+1}/{total_images} pages)")
        
        # Clear the progress message when done
        progress_placeholder.empty()
        return full_text
    except Exception as e:
        st.error(f"Error processing PDF with OCR: {str(e)}")
        st.info("If this is a scanned document, please make sure Tesseract OCR is properly installed.")
        return ""

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            progress_placeholder = st.empty()
            
            progress_placeholder.text(f"Processing PDF with {total_pages} pages...")
            
            for i, page in enumerate(pdf.pages):
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
                
                # Update progress - safely calculate percentage
                percentage = int(100 * (i + 1) / total_pages)
                progress_placeholder.text(f"Processing PDF: {percentage}% ({i+1}/{total_pages} pages)")
            
            # Clear the progress message when done
            progress_placeholder.empty()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
    return text

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def get_document_text(file_path, file_type):
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
        if len(text.strip()) < 50:  # If text is too small, assume it's a scanned PDF
            st.info("PDF appears to be scanned. Using OCR...")
            return extract_text_with_ocr(file_path)
        else:
            st.info("PDF has selectable text. Extracting directly...")
            return text
    elif file_type == "docx":
        st.info("Processing DOCX file...")
        return extract_text_from_docx(file_path)
    return ""

def extract_questions(text):
    question_pattern = r"([A-Z][^?.!]*\?)"
    questions = re.findall(question_pattern, text)
    filtered_questions = [q.strip() for q in questions if len(q.split()) > 2]
    return filtered_questions

def get_gpt_answer(question, openai):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a domain expert providing factual, concise answers. Never mention AI, LLMs, or language models in your responses. Never say 'As an AI' or similar phrases. Respond in a natural, human-like manner with factual information only."},
                {"role": "user", "content": question}
            ],
            max_tokens=150
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error generating answer: {str(e)}")
        return "Unable to generate an answer. Please check your API key and try again."

def save_to_docx(questions_answers, output_file):
    doc = Document()
    doc.add_heading("Questions & Answers from Document", level=1)
    for i, (question, answer) in enumerate(questions_answers, start=1):
        doc.add_paragraph(f"Question {i}: {question}", style="Heading 2")
        doc.add_paragraph(f"{answer}\n")
    doc.save(output_file)
    return output_file

def validate_api_key(api_key):
    if not api_key or len(api_key) < 20:
        return False
    
    # Simple validation for key format
    if not (api_key.startswith('sk-') or api_key.startswith('org-')):
        return False
        
    return True

def get_api_key():
    # First check environment variable
    api_key = os.getenv("OPENAI_API_KEY")
    
    # Next check session state
    if not validate_api_key(api_key) and st.session_state.api_key:
        api_key = st.session_state.api_key
        
    # Finally check config file (imported here to avoid circular imports)
    if not validate_api_key(api_key):
        try:
            import config
            api_key = config.OPENAI_API_KEY
        except (ImportError, AttributeError):
            pass
            
    return api_key

def main():
    # Initialize the sidebar with setup information
    st.sidebar.title("Setup Information")
    
    # Check OpenAI dependency
    openai = check_openai_version()
    
    # Setup Tesseract
    tesseract_available = setup_tesseract()
    
    # Main application
    st.title("📄 Document Question Extractor and Answer Generator")
    st.write("Upload a PDF or DOCX file to extract questions and generate expert answers.")

    # Get API key from environment or config
    current_api_key = get_api_key()
    
    # API Key input
    api_key_input = st.text_input(
        "Enter your OpenAI API Key:",
        value="" if not validate_api_key(current_api_key) else "",
        type="password", 
        placeholder="sk-..." if not validate_api_key(current_api_key) else "Using saved API key",
        help="Your API key will be used only for this session and won't be stored permanently."
    )
    
    # Update session state if user provided a key
    if api_key_input:
        if validate_api_key(api_key_input):
            st.session_state.api_key = api_key_input
            current_api_key = api_key_input
            st.success("API key set for this session!")
        else:
            st.error("Invalid API key format. Please check your key.")
    
    # API key status
    if validate_api_key(current_api_key):
        st.sidebar.success("✅ API Key configured")
        openai.api_key = current_api_key
    else:
        st.warning("Please enter your OpenAI API key to proceed.")
        st.sidebar.error("❌ API Key not configured")
        openai.api_key = None
    
    # Create columns for the file uploader and status display
    col1, col2 = st.columns([2, 1])
    
    # File uploader
    with col1:
        uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx"])
    
    if uploaded_file and validate_api_key(current_api_key):
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        # Check if this is a PDF that might need OCR but tesseract is not available
        if file_type == "pdf" and not tesseract_available:
            st.warning("⚠️ Note: Scanned PDFs cannot be processed because Tesseract OCR is not installed. Text-based PDFs will still work.")
            
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            temp_file_path = tmp_file.name
        
        try:
            # Display status in the second column
            with col2:
                st.info(f"📄 File: {uploaded_file.name}")
            
            with st.spinner("Processing document..."):
                # Extract questions from document
                full_text = get_document_text(temp_file_path, file_type)
                questions = extract_questions(full_text)
                st.session_state.questions = questions
                
                if not questions:
                    st.error("❌ No questions were extracted from the document.")
                else:
                    # Create columns for count display and generate button
                    col1, col2 = st.columns([3, 1])
                    
                    # Display only the number of questions extracted in the first column
                    with col1:
                        st.success(f"✅ Successfully extracted {len(questions)} questions from the document!")
                    
                    # Generate answers button in the second column - more prominent placement
                    with col2:
                        generate_button = st.button("Generate Expert Answers", type="primary", use_container_width=True)
                    
                    # Show a small sample of questions (just first 3) as a preview
                    with st.expander("Preview of extracted questions"):
                        for i, q in enumerate(questions[:3], 1):
                            st.write(f"**Q{i}:** {q}")
                        if len(questions) > 3:
                            st.write(f"_...and {len(questions)-3} more questions_")
                    
                    # Generate answers when button is clicked
                    if generate_button:
                        with st.spinner("Generating answers..."):
                            questions_answers = []
                            progress_placeholder = st.empty()
                            status_text = st.empty()
                            
                            total_questions = len(questions)
                            for i, question in enumerate(questions):
                                status_text.text(f"Generating answer for question {i+1} of {total_questions}")
                                percentage = int(100 * (i + 1) / total_questions)
                                progress_placeholder.progress(percentage / 100)
                                
                                answer = get_gpt_answer(question, openai)
                                questions_answers.append((question, answer))
                            
                            status_text.empty()
                            progress_placeholder.empty()
                            
                            # Store in session state
                            st.session_state.questions_answers = questions_answers
                            
                            # Save to DOCX
                            output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
                            save_to_docx(questions_answers, output_file)
                            
                            # Display download button
                            st.success("✅ All answers generated successfully!")
                            with open(output_file, "rb") as file:
                                btn = st.download_button(
                                    label="Download Q&A Document",
                                    data=file,
                                    file_name="Document_QA.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            
                            # Display questions and answers together
                            st.subheader("Generated Questions & Answers")
                            
                            # Create expandable sections for each Q&A pair
                            for i, (q, a) in enumerate(questions_answers, 1):
                                with st.expander(f"Question {i}: {q[:80]}{'...' if len(q) > 80 else ''}"):
                                    st.markdown(f"**Full Question:** {q}")
                                    st.markdown(f"**Answer:** {a}")
            
            # If we have answers from a previous run, keep them displayed
            if not generate_button and st.session_state.questions_answers:
                st.subheader("Previously Generated Answers")
                # Provide option to download previous results
                if 'previous_output_file' not in st.session_state:
                    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
                    save_to_docx(st.session_state.questions_answers, output_file)
                    st.session_state.previous_output_file = output_file
                
                with open(st.session_state.previous_output_file, "rb") as file:
                    st.download_button(
                        label="Download Previous Q&A Document",
                        data=file,
                        file_name="Document_QA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                for i, (q, a) in enumerate(st.session_state.questions_answers, 1):
                    with st.expander(f"Question {i}: {q[:80]}{'...' if len(q) > 80 else ''}"):
                        st.markdown(f"**Full Question:** {q}")
                        st.markdown(f"**Answer:** {a}")
                
        finally:
            # Clean up the temp file
            try:
                os.unlink(temp_file_path) 
            except:
                pass

    # API Key instructions
    with st.sidebar.expander("How to get an API Key"):
        st.markdown("""
        1. Go to [OpenAI's website](https://platform.openai.com/signup)
        2. Create an account or sign in
        3. Navigate to API keys section
        4. Create a new API key
        5. Copy and paste it in the field above
        """)

    # Footer with hidden developer contact
    st.markdown("---")
    footer_html = '''
    <div style="text-align: center; margin-top: 1em; margin-bottom: 1em;">
        Created with ❤️ | Upload documents to extract questions and generate expert answers
        <span style="display: none;">
            <a href="mailto:umerfarok.dev@gmail.com">Contact Developer</a>
        </span>
        <div style="font-size: 0.7em; color: #888; margin-top: 0.5em;">
            <a href="#" onclick="window.location.href='mailto:umerfarok.dev@gmail.com';" 
               style="text-decoration: none; color: inherit;">Contact Support</a>
        </div>
    </div>
    ''' 
    st.markdown(footer_html, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
